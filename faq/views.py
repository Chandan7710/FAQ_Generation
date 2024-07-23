from django.conf import settings
from langchain.embeddings import HuggingFaceEmbeddings
from llama_index.embeddings.langchain import LangchainEmbedding
from llama_index.llms.together import TogetherLLM
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader
from llama_index.core import StorageContext
from llama_index.core.postprocessor import SentenceTransformerRerank
from llama_index.core.extractors import TitleExtractor, QuestionsAnsweredExtractor
from llama_index.core.node_parser import TokenTextSplitter
from llama_index.core.ingestion import IngestionPipeline
from django.shortcuts import render
from copy import deepcopy
import qdrant_client
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader
import qdrant_client
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core import ChatPromptTemplate
from llama_index.core.llms import ChatMessage, MessageRole
from threading import Lock
from docx import Document
from django.http import HttpResponse
from django.template.loader import render_to_string
import os
import re
from django.http import FileResponse, Http404

# Initializing reraking model
rerank = SentenceTransformerRerank(model="cross-encoder/ms-marco-MiniLM-L-2-v2", top_n=7)

# Setting the message role for system and user using prompt
chat_text_qa_msgs = [
    ChatMessage(
        role=MessageRole.SYSTEM,
        content=(
            """You are an e-government online assistant chatbot system specifically developed for e-Governance Policy Initiatives under Digital India. 
            Your goal is to answer questions as accurately as possible based on the instructions and context provided."""
            " If the question is not related to the uploaded document, respond with 'I can only answer questions related to e-Governance Policy Initiatives under Digital India.'"
            " For general questions like 'Hi', 'How are you?', or 'Who are you?', respond accordingly, mentioning that you're here to assist with e-government policy using the context provided in the document and how you can help the user."
            " Given the context information and not prior knowledge."
            " Make sure to look for headings, subheadings, and key terms that match the question context."

        ),
    ),
    ChatMessage(
        role=MessageRole.USER,
        content=(
            "Context information is below.\n"
            "---------------------\n"
            "{context_str}\n"
            "---------------------\n"
            "Given the context information and not prior knowledge, "

            "answer the question: {query_str} provided in bullet points or numbered list where appropriate.\n"

        ),
    ),
]

text_qa_template = ChatPromptTemplate(chat_text_qa_msgs)
def e_faq(request):
    # return render(request, 'doc_chat.html')
    return render(request, 'faq.html')


# Global variables
index = None
title_questions_dict = None
initialization_lock = Lock()


def initialize_index_and_questions():
    """This view index the documents and extracting the question and answer and storing that in a dictionary and storing the document embeddings in to the QdrantVectorStore"""
    global index, title_questions_dict
    # Checks for initialization: It verifies if the global variables `index` and `title_questions_dict` have already been initialized.
    if index is None or title_questions_dict is None:
        # If not initialized, it acquires a lock using `initialization_lock` to ensure thread-safe initialization.
        with initialization_lock:
            if index is None or title_questions_dict is None:
                # Loads documents from a directory.
                documents = SimpleDirectoryReader("./e_document").load_data()

                # Initializes various models for embedding documents, using a large language model, and processing documents (splitting text, extracting titles and questions).
                embed_model = LangchainEmbedding(HuggingFaceEmbeddings(
                    model_name=settings.SENTENCE_EMBEDDING_MODEL))
                settings.embed_model = embed_model

                # Initialize the LLM model(Using Llama3 model through Together API)
                llm = TogetherLLM(
                    model=settings.LLM_MODEL, api_key=settings.LLM_API_KEY)
                settings.llm = llm

                # Set context window size
                settings.context_window = settings.CONTEXT_WINDOW_SIZE

                # TokenTextSplitter helps with splitting text based on word tokens.
                text_splitter = TokenTextSplitter(
                    separator=" ", chunk_size=1000, chunk_overlap=64)
                # TitleExtractor helps extract titles from documents or text passages.
                title_extractor = TitleExtractor(nodes=30)
                # QuestionsAnsweredExtractor helps in extracting potential questions and answer from the section.
                qa_extractor = QuestionsAnsweredExtractor(questions=3)

                # Processes the loaded documents using the pipeline
                pipeline = IngestionPipeline(
                    transformations=[text_splitter, title_extractor, qa_extractor])
                nodes = pipeline.run(documents=documents,
                                     in_place=True, show_progress=True)
                title_questions_dict = {}
                # Iterates through processed documents and builds a dictionary mapping titles to associated questions and summaries.
                for node in nodes:
                    title = node.metadata.get('document_title', '').strip()
                    summary = node.metadata.get('summary', '').strip()
                    questions = node.metadata.get(
                        'questions_this_excerpt_can_answer', '').strip()
                    # Filters documents based on title
                    # Check if the title starts with "e-kranti" or comes after it alphabetically
                    if title.lower() >= 'e-kranti':
                        # Extract questions from the formatted string
                        if questions:
                            # Extracts questions from a formatted string within each document.
                            extracted_questions = [line.strip() for line in questions.split('\n') if line.strip(
                            ).startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.')]

                            # Add the extracted questions to the title_questions_dict
                            if title not in title_questions_dict:
                                title_questions_dict[title] = {
                                    'questions': [], 'summary': summary}

                            title_questions_dict[title]['questions'].extend(
                                extracted_questions)

                # Initialize the Qdrant vector store
                client = qdrant_client.QdrantClient(path="qdrant_faq")
                # Creates a connection to a vector store for storing document embeddings
                vector_store = QdrantVectorStore(
                    client=client, collection_name="text_collection")
                storage_context = StorageContext.from_defaults(
                    vector_store=vector_store)
                # Creates an index object likely using the processed documents (without metadata) and the vector store connection for efficient document retrieval.
                nodes_no_metadata = deepcopy(nodes)
                index = VectorStoreIndex(
                    nodes=nodes_no_metadata, storage_context=storage_context)
                # print("Parsing is done in FAQ")


initialize_index_and_questions()


def filter_response(response_text):
    unwanted_phrases = [
        'sure', 'I am happy to help you', 'sure I will help you']
    for phrase in unwanted_phrases:
        response_text = response_text.replace(phrase, '')
    return response_text


def format_answer(answer):
    """This function converts the raw answer text into a more user-friendly format with bullet points, 
        numbered lists, and potentially bold text for emphasis, removing unnecessary characters using regular expression"""
    # Remove asterisks used for bold formatting
    answer = re.sub(r'\*\*', '', answer)

    lines = answer.split('\n')
    formatted_lines = []
    in_list = False
    list_type = None

    for line in lines:
        # Check for numbered list
        numbered_match = re.match(r'^\d+\.\s(.+)', line)
        # Check for asterisk list
        asterisk_match = re.match(r'^\*\s(.+)', line)

        if numbered_match:
            if not in_list or list_type != 'ol':
                if in_list:  # Close the previous list
                    formatted_lines.append('</ol>')
                formatted_lines.append('<ol>')
                in_list = True
                list_type = 'ol'
            formatted_lines.append(
                f'<li>{numbered_match.group(1).strip()}</li>')

        elif asterisk_match:
            if not in_list or list_type != 'ul':
                if in_list:  # Close the previous list
                    formatted_lines.append('</ul>')
                formatted_lines.append('<ul>')
                in_list = True
                list_type = 'ul'
            formatted_lines.append(
                f'<li>{asterisk_match.group(1).strip()}</li>')

        else:
            if in_list:  # Close the previous list
                formatted_lines.append(
                    '</ul>' if list_type == 'ul' else '</ol>')
                in_list = False
            # Wrap non-list lines in paragraphs or handle them appropriately
            formatted_lines.append(f'<p>{line.strip()}</p>')

    # Close any open list tags
    if in_list:
        formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')

    # Combine all formatted lines
    formatted_output = ''.join(formatted_lines)

    return formatted_output


def generate_faq(request):
    """This function takes a request object and generates a Frequently Asked Questions (FAQ) document in both Word (.docx) and HTML formats."""
    global index, text_qa_template, title_questions_dict  # Corrected index name

    # Define the output file path
    output_file = os.path.join(settings.BASE_DIR, 'faq_document.docx')

    # Create a new Document
    doc = Document()
    faq_data = []

    # Iterate through the title_questions_dict
    for full_title, data in title_questions_dict.items():
        # Extract the title and summary
        if '\n' in full_title:
            title, summary = full_title.split('\n', 1)
        else:
            title = full_title
            summary = ""

        # Remove unwanted introductory text and format the title and overview correctly
        title = title.replace(
            "The comprehensive title for this document is indeed:", "").strip()
        summary = summary.replace("Summary:", "").strip()

        # Prepare the FAQ data
        faq_item = {"section": title, "overview": summary, "questions": []}

        # Write the title to the document
        doc.add_heading(title.strip(), level=1)

        # Write the summary to the document
        if summary.strip():
            doc.add_heading('Summary', level=2)
            doc.add_paragraph(summary.strip())

        # Process the FAQ section
        if data['questions']:
            doc.add_heading('FAQ', level=2)

            for question in data['questions']:
                # Query the index for the answer
                response = index.as_query_engine(
                    text_qa_template=text_qa_template, similarity_top_k=3, node_postprocessors=[rerank]).query(question)

                # Extract the text from the response object and remove initial content
                response_text = str(response).split('Answer:', 1)[-1].strip()
                filtered_response = filter_response(response_text)

                # Format the answer for HTML
                formatted_response = format_answer(filtered_response)

                # Append the question and filtered response to the FAQ item
                faq_item["questions"].append(
                    {"question": question, "answer": formatted_response})

                # Write the question and response to the document
                doc.add_heading('Question:', level=3)
                doc.add_paragraph(question, style='List Bullet')

                doc.add_heading('Answer:', level=3)
                doc.add_paragraph(filtered_response)

                # print(f"Processed question for title '{title}': {question}")

            doc.add_paragraph()  # Add space between sections
            doc.add_page_break()  # Add page break between titles

        # Append the FAQ item to the FAQ data
        faq_data.append(faq_item)

    # Save the document
    doc.save(output_file)
    # print(f"Results saved in file: {output_file}")

    # Render the FAQ data in the HTML template
    return render(request, 'faq.html', {'faq_data': faq_data})


def download_file(request):
    """This function handles downloading the generated FAQ document created by the generate_faq function."""
    # Define the path to the file you want to serve
    file_path = os.path.join(settings.BASE_DIR, 'faq_document.docx')

    # Check if the file exists
    if not os.path.exists(file_path):
        raise Http404("File does not exist")

    # Open the file for reading content
    file_handle = open(file_path, 'rb')

    # Create the response object with the appropriate content type
    response = FileResponse(
        file_handle, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="faq_document.docx"'

    return response
